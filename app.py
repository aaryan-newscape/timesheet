import os
from flask import Flask, render_template, request, redirect, url_for
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime, timedelta
import pandas as pd

app = Flask(__name__)

def set_table_border(table):
    tbl = table._element
    tblPr = tbl.xpath(".//w:tblPr")[0]
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def get_friday_date_of_current_week(date):
    day_of_week = date.weekday()
    friday = date + timedelta(days=(4 - day_of_week))
    return friday.strftime('%m-%d-%Y')

@app.route('/', methods=['GET', 'POST'])
def form():
    if request.method == 'POST':
        employee_name = request.form['employee_name']
        project_name = request.form['project_name']

        # Validate employee name format
        name_parts = employee_name.split()
        if len(name_parts) < 2:
            return "Please enter a valid employee name in the format 'First Last'.", 400
        name = name_parts[0]
        surname = name_parts[-1]

        # Time Log Data
        dates = request.form.getlist('date')
        start_times = request.form.getlist('start_time')
        end_times = request.form.getlist('end_time')
        regular_hours = request.form.getlist('regular_hours')
        holiday_hours = request.form.getlist('holiday_hours')
        total_hours = request.form.getlist('total_hours')

        if not all(dates) or not all(start_times) or not all(end_times) or not all(regular_hours) or not all(holiday_hours) or not all(total_hours):
            return "All fields in the Time Log are required.", 400

        time_log_data = {
            "Date": dates,
            "Start Time": start_times,
            "End Time": end_times,
            "Regular Hours": regular_hours,
            "Holiday, PTO, OOO": holiday_hours,
            "Total Hours": total_hours
        }

        # Calculate weekly totals
        total_regular_hours = sum([float(hour) for hour in regular_hours if hour])
        total_holiday_hours = sum([float(hour) for hour in holiday_hours if hour])
        total_all_hours = total_regular_hours + total_holiday_hours

        # Task Log Data
        tasks = request.form.getlist('task')
        details = request.form.getlist('detail')
        hours = request.form.getlist('hours')
        statuses = [request.form.get(f'status-{i}') if request.form.get(f'status-{i}') else '' for i in range(len(tasks))]

        if not all(tasks) or not all(details) or not all(hours) or not all(statuses):
            return "All fields in the Task Log are required.", 400

        task_log_data = {
            "Tasks": tasks,
            "Details": details,
            "Hours": hours,
            "Status": statuses
        }

        # Validate total hours
        if total_all_hours < 40:
            return "The sum of Regular Hours and Holiday, PTO, OOO hours must be at least 40.", 400

        # Validate task log total hours
        task_log_total_hours = sum([float(hour) for hour in hours if hour])
        if total_regular_hours != task_log_total_hours:
            return "The total hours in Task Log must match the total Regular Hours.", 400

        # Create DataFrames
        time_log_df = pd.DataFrame(time_log_data)
        task_log_df = pd.DataFrame(task_log_data)

        # Create Word Document
        document = Document()
        document.add_heading('Employee Time Log and Task Log', 0)
        document.add_paragraph(f"EMPLOYEE NAME: {employee_name}")
        document.add_paragraph(f"PROJECT NAME: {project_name}")
        document.add_paragraph("\n")

        # Add Time Log Table
        document.add_heading('Time Log', level=1)
        table = document.add_table(rows=1, cols=len(time_log_df.columns))
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(time_log_df.columns):
            hdr_cells[i].text = column

        for index, row in time_log_df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

        # Add Weekly Totals Row
        row_cells = table.add_row().cells
        row_cells[0].text = "WEEKLY TOTALS"
        row_cells[1].text = ""
        row_cells[2].text = ""
        row_cells[3].text = str(total_regular_hours)
        row_cells[4].text = str(total_holiday_hours)
        row_cells[5].text = str(total_all_hours)

        # Set table borders
        set_table_border(table)

        # Add Task Log Table
        document.add_heading('Task Log', level=1)
        table = document.add_table(rows=1, cols=len(task_log_df.columns))
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(task_log_df.columns):
            hdr_cells[i].text = column

        for index, row in task_log_df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

        # Set table borders
        set_table_border(table)

        # Get current date and calculate the Friday of the current week
        current_date = datetime.now()
        friday_date = get_friday_date_of_current_week(current_date)

        # Create folder for the current week if it doesn't exist
        folder_name = f"Timesheet_{friday_date}"
        if not os.path.exists(folder_name):
            os.makedirs(folder_name)

        # Save Document with specified format
        document_filename = os.path.join(folder_name, f"Timesheet_{name}_{surname}_{friday_date}.docx")
        document.save(document_filename)

        return redirect(url_for('form'))

    return render_template('form.html')

if __name__ == '__main__':
    app.run(debug=True)
